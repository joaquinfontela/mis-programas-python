from typing import Text
from docx import Document


def main():

    show_instructions()
    path = ask_for_path()
    doc = Document(path)
    orig_seq, new_seq = ask_words()
    for p in doc.paragraphs:
        paragraph_text = p.text
        p.text = paragraph_text.replace(orig_seq, new_seq)
    doc.save(f"new-{path}")
    return


def show_instructions():
    print("The following program will replace each appearance of a character sequence in a .docx file with a new one in a new file. The old file will remain with no modifications. If its name was 'X.docx', the new file name will be 'new-X.docx'.")
    print("Firstly, it will ask you for the file name: it MUST be a .docx (not .doc) extension file and, also the submitted path MUST be on the same folder as the one where you are running this program. You can enter the file name with its extension (sample.docx) or only the name (sample).")
    input("Press ENTER to continue\n")


def ask_for_path():
    path = input("Enter the exact file name: ")
    if path[-5:] != '.docx':
        path += '.docx'
    try:
        doc = Document(path)
    except:
        print(
            f"FATAL ERROR: The submitted path ({path}) is not a valid .docx document.")
        exit()
    return path


def ask_words():
    original_sequence = input(
        "Enter the exact character sequence you want to be replaced: ")
    new_sequence = input(
        "Enter the exact character sequence you want to replace the old one: ")
    return original_sequence, new_sequence


if __name__ == "__main__":
    main()
